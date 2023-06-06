import os
import win32com.client as win32
from docx import Document
from paragraphs_to_delete import paragraphs_to_delete


def delete_paragraph_by_content(doc, target_content):
    # Store paragraph text, not objects
    for p in doc.paragraphs:
        if target_content in p.text:
            paragraphs_to_delete.append(p.text)  # Append paragraph text

    for text in paragraphs_to_delete:
        for p in doc.paragraphs:
            if text == p.text:  # Compare paragraph text
                p.clear()


def delete_empty_bullet_points(doc):
    paragraphs_to_delete = []

    for paragraph in doc.paragraphs:
        if paragraph.style.name == "pf1" and not paragraph.text:
            paragraphs_to_delete.append(paragraph)

    for paragraph in paragraphs_to_delete:
        p = paragraph._element
        p.getparent().remove(p)


def delete_paragraphs_under_header(document, header_name):
    header_found = False

    for i, paragraph in enumerate(document.paragraphs):
        if paragraph.style.name == "Heading 2" and paragraph.text == header_name:
            header_found = True
        elif header_found:
            # Delete paragraphs under the header until next header or end of document
            document._body._body.remove(paragraph._element)
            paragraph._element.clear_content()

            # Adjust the index after removing the paragraph
            i -= 1

            # Check if the next paragraph is a header to stop deleting
            if (
                i + 1 < len(document.paragraphs)
                and document.paragraphs[i + 1].style.name == "Heading 1"
            ):
                break


def run_word_macro(modified_file_name):
    # Create an instance of the Word application
    word_app = win32.gencache.EnsureDispatch("Word.Application")
    modified_file_path = os.path.join(os.getcwd(), modified_file_name)

    print(modified_file_name)
    try:
        # Open the Word document
        doc = word_app.Documents.Open(modified_file_path)

        # Run the macro
        word_app.Application.Run("DeleteFirstImageAndTableOfContents")

        # Save and close the document
        doc.Save()
        doc.Close()
    except Exception as e:
        print("Error:", e)
    finally:
        # Quit the Word application
        word_app.Quit()


def main():
    # Get the current directory of the script
    script_directory = os.path.dirname(os.path.abspath(__file__))
    # Change the current working directory
    os.chdir(script_directory)

    original_doc_filename = ""
    orginal_doc_path = ""
    modified_doc_path = ""

    while True:
        # Load the document
        original_doc_filename = input("Enter the filename: ")
        orginal_doc_path = os.path.join("word_doc", original_doc_filename + ".docx")

        if os.path.exists(orginal_doc_path):
            break
        else:
            print("File does not exist. Please try again")

    modified_doc_path = os.path.splitext(orginal_doc_path)[0] + "_modified.docx"
    doc = Document(orginal_doc_path)

    # Specify the header name to delete everything under
    header_name = "Quick Resources"
    # Call the function to delete paragraphs under the specified header
    delete_paragraphs_under_header(doc, header_name)

    for target_content in paragraphs_to_delete:
        delete_paragraph_by_content(doc, target_content)

    delete_empty_bullet_points(doc)

    # Save the modified document
    doc.save(modified_doc_path)

    # Call the function to run the Word macro
    run_word_macro(modified_doc_path)

    os.startfile(modified_doc_path)


if __name__ == "__main__":
    main()
