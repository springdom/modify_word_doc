import os
from docx import Document


def print_document_info(document_path):
    # Create a Document object from the specified document path
    document = Document(document_path)

    # Print general document information
    print(f"Number of paragraphs: {len(document.paragraphs)}")
    print(f"Number of tables: {len(document.tables)}")
    print(f"Number of sections: {len(document.sections)}")

    # Print paragraph details
    print("Paragraphs:")
    for i, paragraph in enumerate(document.paragraphs, start=1):
        print(f"Paragraph {i}:")
        print(f"Text: {paragraph.text}")  # Print the text content of the paragraph
        print(
            f"Style: {paragraph.style.name}"
        )  # Print the name of the paragraph's style
        print()

    # Print table details
    print("Tables:")
    for i, table in enumerate(document.tables, start=1):
        print(f"Table {i}:")
        print(
            f"Number of rows: {len(table.rows)}"
        )  # Print the number of rows in the table
        print(
            f"Number of columns: {len(table.columns)}"
        )  # Print the number of columns in the table
        print()

    # Print section details
    print("Sections:")
    for i, section in enumerate(document.sections, start=1):
        print(f"Section {i}:")
        print(f"Header: {section.header}")  # Print the content of the section's header
        print(f"Footer: {section.footer}")  # Print the content of the section's footer
        print()


def main():
    # Path to the document file
    script_directory = os.path.dirname(os.path.abspath(__file__))
    os.chdir(script_directory)
    doc_file_name = ""
    doc_file_path = ""

    while True:
        doc_file_name = input("Enter Filename: ")
        doc_file_path = os.path.join("word_doc", doc_file_name + ".docx")
        print(doc_file_path)
        
        if os.path.exists(doc_file_path):
            break
        else:
            print("File does not exist. Please try again")

    # Call the function to print the document information
    print_document_info(doc_file_path)


if __name__ == "__main__":
    main()
